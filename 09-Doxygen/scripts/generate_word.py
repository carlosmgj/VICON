#!/usr/bin/env python3
"""
generate_word.py
================
Genera un documento Word con la documentación del diseño VHDL.

Uso:
    py -3.11 scripts/generate_word.py <project_path>
                [--config doc_config.yaml]
                [--out docs/design_document.docx]
                [--diagrams-dir docs/diagrams]
                [--plantuml plantuml]
                [--top nombre_modulo_top]

Pasos internos:
    1. Parsea todos los .vhd con GhdlParser (de vhdl_dox_filter)
    2. Extrae jerarquía de instanciaciones (de generate_hierarchy)
    3. Genera diagramas PlantUML PNG (entidad + FSM por módulo)
    4. Construye module_data.json
    5. Llama a scripts/word_generator.js para producir el .docx
"""

import sys
import json
import argparse
import subprocess
from pathlib import Path
from datetime import datetime

# Añadir carpeta de scripts al sys.path para importar módulos existentes
# Directorio donde está este script (siempre absoluto, independiente del CWD)
SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))

from vhdl_dox_filter import GhdlParser, VhdlEntity
from generate_hierarchy import HierarchyParser, VhdlModule

try:
    import yaml
    _HAS_YAML = True
except ImportError:
    _HAS_YAML = False


# =============================================================================
# GENERADORES DE PLANTUML
# =============================================================================

class PumlEntityGenerator:
    """Genera texto PlantUML para el diagrama de E/S de una entidad VHDL."""

    def generate(self, entity: VhdlEntity) -> str:
        inputs  = [p for p in entity.ports if p.direction in ('in', 'inout')]
        outputs = [p for p in entity.ports if p.direction in ('out', 'buffer', 'inout')]

        lines = [
            '@startuml',
            'skinparam componentStyle rectangle',
            'skinparam defaultFontName Helvetica',
            'skinparam component {',
            '  BackgroundColor #E8F4FD',
            '  BorderColor #2E86AB',
            '  FontColor #1a1a1a',
            '}',
            'skinparam arrow {',
            '  Color #2E86AB',
            '}',
            'left to right direction',
            '',
        ]
        for p in inputs:
            type_clean = p.type_str.replace('(', '[').replace(')', ']')
            lines.append(f'() "{p.name}\\n[{type_clean}]" as i_{p.name}')
        for p in outputs:
            type_clean = p.type_str.replace('(', '[').replace(')', ']')
            lines.append(f'() "{p.name}\\n[{type_clean}]" as o_{p.name}')
        lines.append(f'\n[{entity.name}] as ENT\n')
        for p in inputs:
            lines.append(f'i_{p.name} --> ENT')
        for p in outputs:
            lines.append(f'ENT --> o_{p.name}')
        lines += ['', '@enduml']
        return '\n'.join(lines)


class PumlFsmGenerator:
    """Genera texto PlantUML para el diagrama FSM de una entidad VHDL."""

    def generate(self, entity: VhdlEntity) -> str:
        initial = entity.enum_types[0].states[0] if entity.enum_types else None
        lines = [
            '@startuml',
            'skinparam state {',
            '  BackgroundColor #E8F4FD',
            '  BorderColor #2E86AB',
            '  FontName Helvetica',
            '}',
            '',
        ]
        if initial:
            lines.append(f'[*] --> {initial}\n')

        all_states = entity.enum_types[0].states if entity.enum_types else []
        for state in all_states:
            lines.append(f'state {state}')
        lines.append('')

        for tr in entity.fsm_transitions:
            label = f' : {tr.condition}' if tr.condition else ''
            lines.append(f'{tr.from_state} --> {tr.to_state}{label}')

        if entity.fsm_pointers:
            lines.append('')
            for ptr in entity.fsm_pointers:
                node = f'PTR_{ptr.name}'
                lines.append(f'state {node} : {ptr.name}\\n(dynamic return)')
                for fs in ptr.used_from:
                    lines.append(f'{fs} --> {node}')
                for val in ptr.possible_values:
                    lines.append(f'{node} --> {val}')

        lines += ['', '@enduml']
        return '\n'.join(lines)


# =============================================================================
# RUNNER DE PLANTUML
# =============================================================================

class PlantUmlRunner:
    """Ejecuta PlantUML para convertir texto .puml a un fichero .png."""

    def __init__(self, plantuml_jar: 'str | None' = None, plantuml_cmd: str = 'plantuml'):
        """
        plantuml_jar — ruta al plantuml.jar → invoca 'java -jar <jar>'
        plantuml_cmd — comando directo si no hay jar (ej. 'plantuml')
        """
        if plantuml_jar:
            self.shell_cmd = f'java -jar "{plantuml_jar}"'
        else:
            self.shell_cmd = plantuml_cmd

    def generate_png(self, puml_text: str, out_png: Path) -> bool:
        """Genera un PNG a partir de texto PlantUML. Devuelve True si OK."""
        out_png.parent.mkdir(parents=True, exist_ok=True)
        puml_path = out_png.with_suffix('.puml')
        puml_path.write_text(puml_text, encoding='utf-8')

        # Ejecutar vía shell=True para que Windows maneje 'java -jar ...' correctamente
        full_cmd = f'{self.shell_cmd} -tpng -charset UTF-8 "{puml_path}"'
        try:
            result = subprocess.run(
                full_cmd, shell=True, capture_output=True, timeout=60
            )
            generated = puml_path.with_suffix('.png')
            if generated.exists() and generated != out_png:
                generated.rename(out_png)
            if not out_png.exists():
                stderr_txt = result.stderr.decode(errors='replace')[:200] if result.stderr else ''
                print(f'[generate_word]   ⚠ PlantUML sin PNG (rc={result.returncode}) {stderr_txt}',
                      file=sys.stderr)
                return False
            return True
        except subprocess.TimeoutExpired:
            print(f'[generate_word] ⚠ PlantUML timeout en {puml_path}', file=sys.stderr)
            return False
        except Exception as e:
            print(f'[generate_word] ⚠ PlantUML error: {e}', file=sys.stderr)
            return False


# =============================================================================
# BUILDER DEL JSON INTERMEDIO
# =============================================================================

class ModuleDataBuilder:
    """
    Combina VhdlEntity (pyGHDL) y VhdlModule (regex) en el JSON que
    consume word_generator.js.
    """

    def __init__(self, diagrams_dir: Path, plantuml_jar: str | None = None, plantuml_cmd: str = 'plantuml'):
        self.diagrams_dir = diagrams_dir
        self.puml_runner  = PlantUmlRunner(plantuml_jar=plantuml_jar, plantuml_cmd=plantuml_cmd)
        self.entity_gen   = PumlEntityGenerator()
        self.fsm_gen      = PumlFsmGenerator()

    # ── Punto de entrada ─────────────────────────────────────────────────────

    def build(
        self,
        entities:   dict[str, VhdlEntity],
        modules:    list[VhdlModule],
        top_module: str | None,
        config:     dict,
    ) -> dict:

        # Detectar top automáticamente si hace falta
        if not top_module or top_module == 'auto':
            all_inst = {i.lower() for m in modules for i in m.instantiates}
            roots    = [m.name for m in modules if m.name.lower() not in all_inst]
            top_module = roots[0] if roots else (modules[0].name if modules else 'unknown')

        module_map = {m.name.lower(): m for m in modules}

        # Serializar cada módulo
        modules_data = {}
        for m in modules:
            entity = entities.get(m.name.lower())
            modules_data[m.name] = self._serialize(m, entity)

        # Jerarquía top-down
        hierarchy = self._hierarchy(top_module, module_map, set())

        # Metadatos del documento
        meta = dict(config.get('document', {}))
        if not meta.get('date'):
            meta['date'] = datetime.now().strftime('%d/%m/%Y')

        return {
            'meta':       meta,
            'top_module': top_module,
            'hierarchy':  hierarchy,
            'modules':    modules_data,
        }

    # ── Serialización de un módulo ───────────────────────────────────────────

    def _serialize(self, module: VhdlModule, entity: VhdlEntity | None) -> dict:
        data: dict = {
            'name':               module.name,
            'file':               module.path.name if module.path else '',
            'architecture':       module.architecture,
            'instantiates':       module.instantiates,
            'entity_diagram_png': None,
            'fsm_diagram_png':    None,
            'ports':              [],
            'generics':           [],
            'signals':            [],
            'constants':          [],
            'fsm_transitions':    [],
        }

        if entity is None:
            return data

        data['ports'] = [
            {'name': p.name, 'direction': p.direction, 'type': p.type_str,
             'default': p.default or '', 'description': p.comment or ''}
            for p in entity.ports
        ]
        data['generics'] = [
            {'name': g.name, 'type': g.type_str,
             'default': g.default or '', 'description': g.comment or ''}
            for g in entity.generics
        ]
        data['signals'] = [
            {'name': s.name, 'type': s.type_str,
             'default': s.default or '', 'description': s.comment or ''}
            for s in entity.signals if 'next' not in s.name.lower()
        ]
        data['constants'] = [
            {'name': c.name, 'type': c.type_str,
             'value': c.value or '', 'description': c.comment or ''}
            for c in entity.constants
        ]
        data['fsm_transitions'] = [
            {'from': t.from_state, 'to': t.to_state, 'condition': t.condition or ''}
            for t in entity.fsm_transitions
        ]

        # Diagramas PlantUML → PNG
        if entity.ports:
            png = self.diagrams_dir / f'{entity.name}_entity.png'
            if self.puml_runner.generate_png(self.entity_gen.generate(entity), png):
                data['entity_diagram_png'] = str(png.resolve())
            else:
                print(f'[generate_word]   ⚠ Sin PNG de entidad para {entity.name}', file=sys.stderr)

        if entity.fsm_transitions:
            png = self.diagrams_dir / f'{entity.name}_fsm.png'
            if self.puml_runner.generate_png(self.fsm_gen.generate(entity), png):
                data['fsm_diagram_png'] = str(png.resolve())

        return data

    # ── Árbol de jerarquía ───────────────────────────────────────────────────

    def _hierarchy(self, name: str, module_map: dict, visited: set) -> dict:
        if name.lower() in visited:
            return {'name': name, 'children': [], 'recursive': True}
        visited = visited | {name.lower()}
        mod = module_map.get(name.lower())
        children = []
        if mod:
            for child_name in mod.instantiates:
                children.append(self._hierarchy(child_name, module_map, visited))
        return {'name': name, 'children': children, 'recursive': False}


# =============================================================================
# HELPERS DE CONFIGURACIÓN
# =============================================================================

def load_config(config_path: Path) -> dict:
    if not config_path.exists():
        print(f'[generate_word] ⚠ Config no encontrado: {config_path}, usando defaults', file=sys.stderr)
        return {}
    text = config_path.read_text(encoding='utf-8')
    if _HAS_YAML:
        return yaml.safe_load(text) or {}
    # Fallback: intentar JSON
    try:
        return json.loads(text)
    except Exception:
        print('[generate_word] ⚠ No se pudo parsear config (instala PyYAML: pip install pyyaml)', file=sys.stderr)
        return {}


def find_top_from_config(config: dict) -> str:
    for sec in config.get('sections', []):
        if sec.get('type') == 'modules':
            return sec.get('top_module', 'auto')
    return 'auto'


# =============================================================================
# ENTRY POINT
# =============================================================================

def main():
    ap = argparse.ArgumentParser(description='Genera documento Word de diseño VHDL')
    ap.add_argument('project_path',    help='Ruta raíz del proyecto VHDL')
    ap.add_argument('--config',        default='doc_config.yaml',           help='Fichero YAML de configuración')
    ap.add_argument('--out',           default='docs/design_document.docx',  help='Fichero .docx de salida')
    ap.add_argument('--diagrams-dir',  default='docs/diagrams',              help='Directorio para PNGs intermedios')
    ap.add_argument('--plantuml',      default='plantuml',                   help='Comando directo de PlantUML')
    ap.add_argument('--plantuml-jar',  default=None,                         help='Ruta al plantuml.jar (recomendado en Windows)')
    ap.add_argument('--top',           default=None,                         help='Módulo top explícito')
    args = ap.parse_args()

    project_path = Path(args.project_path)
    if not project_path.exists():
        print(f'[generate_word] Error: {project_path} no existe', file=sys.stderr)
        sys.exit(1)

    config       = load_config(Path(args.config))
    diagrams_dir = Path(args.diagrams_dir)
    diagrams_dir.mkdir(parents=True, exist_ok=True)

    # ── Recoger ficheros VHDL ─────────────────────────────────────────────
    EXCLUDE_DIRS = {'02-IPs', 'ip', '__pycache__', 'node_modules'}
    vhd_files = sorted([
        f for f in list(project_path.rglob('*.vhd')) + list(project_path.rglob('*.vhdl'))
        if not any(part in EXCLUDE_DIRS for part in f.parts)
    ])
    if not vhd_files:
        print(f'[generate_word] No se encontraron ficheros VHDL en {project_path}', file=sys.stderr)
        sys.exit(1)
    print(f'[generate_word] {len(vhd_files)} ficheros VHDL encontrados', file=sys.stderr)

    # ── 1. Parsear entidades con GhdlParser ───────────────────────────────
    ghdl_parser = GhdlParser()
    entities: dict[str, VhdlEntity] = {}
    for vhd in vhd_files:
        try:
            raw = vhd.read_bytes()
            if raw[:2] in (b'\xff\xfe', b'\xfe\xff'):
                src = raw.decode('utf-16')
            elif raw[:3] == b'\xef\xbb\xbf':
                src = raw.decode('utf-8-sig')
            else:
                src = raw.decode('utf-8', errors='replace')
            entity = ghdl_parser.parse(vhd, src)
            if entity:
                entities[entity.name.lower()] = entity
                print(f'[generate_word]   ✓ {vhd.name}: {entity.name} '
                      f'({len(entity.ports)}p {len(entity.signals)}s {len(entity.fsm_transitions)}t)',
                      file=sys.stderr)
        except ImportError:
            print('[generate_word] ⚠ pyGHDL no disponible — se omiten datos detallados de puertos/FSM',
                  file=sys.stderr)
            break
        except Exception as e:
            print(f'[generate_word]   ⚠ {vhd.name}: {e}', file=sys.stderr)

    # ── 2. Parsear jerarquía con HierarchyParser ──────────────────────────
    hier_parser = HierarchyParser()
    all_modules: list[VhdlModule] = []
    for vhd in vhd_files:
        all_modules.extend(hier_parser.parse_file(vhd))
    if not all_modules:
        print('[generate_word] No se encontraron módulos VHDL', file=sys.stderr)
        sys.exit(1)

    # ── 3-4. Generar diagramas + JSON ─────────────────────────────────────
    top_module = args.top or find_top_from_config(config)
    print(f'[generate_word] Módulo top: {top_module}', file=sys.stderr)
    print(f'[generate_word] Generando diagramas PlantUML...', file=sys.stderr)

    builder     = ModuleDataBuilder(diagrams_dir,
                                    plantuml_jar=getattr(args, 'plantuml_jar', None),
                                    plantuml_cmd=args.plantuml)
    module_data = builder.build(entities, all_modules, top_module, config)

    json_path = diagrams_dir / 'module_data.json'
    json_path.write_text(json.dumps(module_data, indent=2, ensure_ascii=False), encoding='utf-8')

    # También guardar config como JSON para que el generador JS no necesite PyYAML
    config_json_path = diagrams_dir / 'doc_config.json'
    config_json_path.write_text(json.dumps(config, indent=2, ensure_ascii=False), encoding='utf-8')

    print(f'[generate_word] JSON listo en {json_path}', file=sys.stderr)

    # ── 5. Llamar a word_generator.js ─────────────────────────────────────
    generator_js = SCRIPT_DIR / 'word_generator.js'
    if not generator_js.exists():
        print(f'[generate_word] Error: no se encuentra {generator_js}', file=sys.stderr)
        sys.exit(1)

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    print(f'[generate_word] Generando Word...', file=sys.stderr)
    result = subprocess.run(
        ['node', str(generator_js),
         str(json_path.resolve()), str(config_json_path.resolve()), str(out_path.resolve())],
        capture_output=True, text=True, encoding='utf-8',
    )

    if result.stdout.strip():
        print(result.stdout.strip(), file=sys.stderr)
    if result.returncode != 0:
        print(f'[generate_word] Error en word_generator.js:\n{result.stderr}', file=sys.stderr)
        sys.exit(1)

    print(f'[generate_word] ✓ Documento generado: {out_path}', file=sys.stderr)

    # ── Validación rápida con python-docx ─────────────────────────────────
    try:
        from docx import Document as DocxDocument
        d = DocxDocument(str(out_path))
        print(f'[generate_word] ✓ Validación OK — {len(d.paragraphs)} párrafos, {len(d.tables)} tablas',
              file=sys.stderr)
    except ImportError:
        pass  # python-docx no instalado, no pasa nada
    except Exception as val_err:
        print(f'[generate_word] ⚠ Documento generado pero python-docx reporta: {val_err}',
              file=sys.stderr)
        print(f'[generate_word]   → Intenta: pip install python-docx --break-system-packages',
              file=sys.stderr)


if __name__ == '__main__':
    main()
